from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "outputs" / "layered-image-representations" / "docs" / "Project_Documentation.docx"
FLOW = Path(__file__).with_name("pipeline_flow.png")
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 98, 110)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margin(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    tbl_pr.append(indent)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            set_cell_margin(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def run_font(run, size=11, color=RGBColor(0, 0, 0), bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        run_font(p.add_run(bold_prefix), bold=True)
        run_font(p.add_run(text[len(bold_prefix):]))
    else:
        run_font(p.add_run(text))
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    run_font(p.add_run(text))
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run_font(p.add_run(text), size={1: 16, 2: 13, 3: 12}[level], color={1: BLUE, 2: BLUE, 3: DARK}[level], bold=True)
    return p


def make_flow_image(path: Path):
    w, h = 1240, 260
    image = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(image)
    try:
        font = ImageFont.truetype("arial.ttf", 23)
        small = ImageFont.truetype("arial.ttf", 17)
    except OSError:
        font = ImageFont.load_default()
        small = font
    boxes = [
        (25, 65, 255, 190, "Single RGB\nimage", "E8EEF5"),
        (335, 65, 565, 190, "Semantic\nsegmentation", "DDEBF7"),
        (645, 65, 875, 190, "Relative depth\nestimation", "DDEBF7"),
        (955, 65, 1215, 190, "RGBA layers +\nmanifest", "EAF2F8"),
    ]
    for x1, y1, x2, y2, text, fill in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=14, fill=f"#{fill}", outline="#2E74B5", width=3)
        lines = text.split("\n")
        for i, line in enumerate(lines):
            bbox = draw.textbbox((0, 0), line, font=font)
            tx = (x1 + x2 - (bbox[2] - bbox[0])) / 2
            ty = 95 + i * 31
            draw.text((tx, ty), line, fill="#0B2545", font=font)
    for x in (275, 585, 895):
        draw.line((x, 127, x + 45, 127), fill="#2E74B5", width=4)
        draw.polygon([(x + 45, 127), (x + 31, 119), (x + 31, 135)], fill="#2E74B5")
    draw.text((25, 215), "Figure 1. Inference pipeline. Layer order is obtained from mean relative depth within each semantic mask.", fill="#59626E", font=small)
    image.save(path)


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for level, size, color, before, after in ((1, 16, BLUE, 16, 8), (2, 13, BLUE, 12, 6), (3, 12, DARK, 8, 4)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_font(footer.add_run("DLCV Project Documentation | 22f2000506"), size=9, color=MUTED)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    make_flow_image(FLOW)
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run_font(title.add_run("LAYERED REPRESENTATIONS FROM A SINGLE IMAGE"), size=23, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run_font(subtitle.add_run("DLCV Project Documentation"), size=14, color=MUTED)
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(16)
    run_font(meta.add_run("Student / GitHub user: "), bold=True)
    run_font(meta.add_run("22f2000506    |    Date: 1 September 2026"))

    add_heading(doc, "Executive summary")
    add_body(doc, "This project transforms a single bitmap RGB image into a re-composable set of transparent layers. A semantic segmentation model identifies interpretable regions, and a monocular-depth model estimates their relative order. The system exports an RGBA image for each layer, a near-to-far manifest, a relative-depth visualization, and a reconstructed composite. It also produces simple albedo and shading proxies as an explicitly labelled stretch feature.")
    doc.add_picture(str(FLOW), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "1. Problem and motivation")
    add_body(doc, "A conventional image merges foreground objects, background material, lighting, and depth into one RGB array. This project investigates whether pretrained vision models can recover an interpretable layered representation from that single observation. Such a representation is useful for object-level editing, basic parallax animation, relighting experiments, and downstream visual analysis.")

    add_heading(doc, "2. Proposed method")
    add_heading(doc, "Semantic grouping", level=2)
    add_body(doc, "Mask2Former produces semantic masks for recognisable scene regions. The expected labels depend on the image but can include person, vehicle, furniture, wall, floor, sky, and vegetation. Small masks below a configurable image-area threshold are ignored to reduce noise.")
    add_heading(doc, "Relative depth ordering", level=2)
    add_body(doc, "Depth Anything V2 estimates dense relative depth from the same image. The mean value inside each retained mask is used to rank the exported layers from near to far. The depth is relative and must not be treated as metric distance.")
    add_heading(doc, "RGBA and appearance outputs", level=2)
    add_body(doc, "For every mask, the original RGB pixels are preserved and the mask becomes the alpha channel. The code also saves a luminance-based shading proxy and a chromatic albedo proxy. These are exploratory visual decompositions, not ground-truth intrinsic-image estimates.")

    add_heading(doc, "3. Deliverables and repository")
    add_body(doc, "The repository contains a command-line runner, the inference pipeline, an offline smoke-test backend, a test, and this documentation. A successful run creates the files listed below.")
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [3000, 6360])
    headers = ["File", "Purpose"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, "F2F4F7")
        p = cell.paragraphs[0]
        run_font(p.add_run(text), bold=True)
    rows = [
        ("01_<label>/rgba.png", "Transparent semantic layer, ordered near to far"),
        ("albedo_proxy.png / shading_proxy.png", "Stretch-feature appearance proxies for the same layer"),
        ("relative_depth.png", "Normalised visualisation of relative monocular depth"),
        ("composite.png", "Image reconstructed by alpha compositing far-to-near layers"),
        ("manifest.json", "Labels, confidence, depth, pixel count, and output locations"),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        run_font(cells[0].paragraphs[0].add_run(left), bold=True)
        run_font(cells[1].paragraphs[0].add_run(right))

    add_heading(doc, "4. How to run")
    add_body(doc, "Install the packages listed in requirements.txt, then run the command below from the repository root.")
    command = doc.add_paragraph()
    command.paragraph_format.left_indent = Inches(0.25)
    command.paragraph_format.space_after = Pt(8)
    run_font(command.add_run("python run.py path\\to\\image.jpg --output output\\scene_01"), size=10, color=DARK)
    add_body(doc, "The first transformer run downloads pretrained weights. An offline `heuristic` backend is available for smoke testing; it segments colour/spatial regions only and should not be presented as the semantic result.")

    add_heading(doc, "5. Evaluation plan")
    add_body(doc, "Evaluate 10-20 images spanning indoor and outdoor scenes, multiple object categories, and visible depth variation. For each test image, include the original image, relative-depth map, layer stack, recomposition, and one albedo/shading-proxy pair in the final submission.")
    for item in (
        "Semantic grouping: visually check whether each meaningful object or background category is assigned to an understandable layer.",
        "Depth ordering: compare manifest order with the scene's apparent near-to-far order.",
        "Recomposition: compare composite.png against the input and note holes, leakage, or boundary errors.",
        "Failure analysis: include at least one difficult case, such as reflections, transparency, severe occlusion, or a small distant object.",
    ):
        add_bullet(doc, item)
    add_body(doc, "Important: do not invent accuracy values. If no manually labelled ground truth is available, use this qualitative protocol and report observations honestly.", bold_prefix="Important: ")

    add_heading(doc, "6. Limitations and responsible use")
    add_body(doc, "Monocular depth is ambiguous and relative. Segmentation models can merge categories, miss small objects, or inherit bias from their training data. Reflective, transparent, and heavily occluded regions are particularly challenging. The layer outputs are for visual-computing experiments and should not be used as the basis for high-stakes decisions.")

    add_heading(doc, "7. Submission evidence to add")
    add_body(doc, "Before submitting, replace this section with 3-5 of your own examples and a concise results table. Keep the link to the public GitHub repository, the exact commands used, and screenshots of the four key outputs.")
    for item in (
        "Input image alongside composite.png",
        "Relative-depth image alongside near-to-far layer thumbnails",
        "One successful scene and one documented failure case",
        "The GitHub repository URL and commit hash used for submission",
    ):
        add_bullet(doc, item)

    add_heading(doc, "References")
    refs = [
        "Cheng, B. et al. Masked-attention Mask Transformer for Universal Image Segmentation. CVPR, 2022.",
        "Yang, L. et al. Depth Anything V2. 2024.",
        "Hugging Face Transformers documentation: image-segmentation and depth-estimation pipelines.",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run_font(p.add_run(ref), size=10)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
